# -*- coding: utf-8 -*-
"""
Created on Mon Jan 20 13:57:41 2020

@author: NFC User
"""


import re
import pandas as pd
import numpy as np
import PyPDF2


from gensim.models import Word2Vec

from win32com import client
from time import strftime

import nltk
from nltk.corpus import stopwords
set(stopwords.words('english'))

from nltk.stem import SnowballStemmer

import os
from os import chdir, getcwd, listdir, path

word = client.DispatchEx("Word.Application")
### to work install  sofficeor libre or antiword... win32com installation
os.chdir('D:/mix_pw/')

### only for converting the doc to pdf
for files in os.listdir(os.getcwd()):
    if files.endswith(".doc"):
            new_name = files.replace(".doc", r".pdf")
            in_file = path.abspath(os.getcwd() + "\\" + files)
            new_file = path.abspath(os.getcwd() + "\\" + new_name)
            doc = word.Documents.Open(in_file)
            print(strftime("%H:%M:%S"), " doc  -> pdf ", path.relpath(new_file))
            doc.SaveAs2(new_file, FileFormat = 17)###file format 16 for docx
            #doc.remove(in_file)
            doc.Close()
entries_pdf=[]
for files in os.listdir(os.getcwd()):
    if files.endswith(".pdf"):
        entries_pdf.append(files)
        
entries_docx=[]
for files in os.listdir(os.getcwd()):
    if files.endswith(".docx"):
        entries_docx.append(files)
        

text=[]
file=[]
for i in range(0,len(entries_pdf)):
    pdfFileObj = open(os.getcwd()+ '\\'+ entries_pdf[i],'rb')
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
    file.append(entries_pdf[i])
    t=''
    for k in range(0,pdfReader.numPages):
        pagobj=pdfReader.getPage(k)
        te=pagobj.extractText()
        t= t +''.join(te)
    text.append(t)
    
import  docx
for i in range(0,len(entries_docx)):
    doc= docx.Document(os.getcwd()+ '\\'+ entries_docx[i])
    file.append(entries_docx[i])
    tee=''
    for para in doc.paragraphs:
        tw=para.text
        tee=tee+''.join(tw)
    text.append(tee)


    
    

def processing(text1):
    # first remove the web links to clean the text
    ## for removing websitelinks
    text1=list(map(lambda x: re.sub(r'''(?i)\b((?:https?://|www\d{0,3}[.]|[a-z0-9.\-]+[.][a-z]{2,4}/)(?:[^\s()<>]+|\(([^\s()<>]+|(\([^\s()<>]+\)))*\))+(?:\(([^\s()<>]+|(\([^\s()<>]+\)))*\)|[^\s`!()\[\]{};:'".,<>?«»“”‘’]))''', " ",x, flags=re.MULTILINE),text1))
    ##for removing mail ids
    text2=list(map(lambda x: re.sub(r'\S+@\S+',' ',x),text1))
    ##for cleaning text which has .
    text3=list(map(lambda x: re.sub(r'[^A-Za-z0-9(),!.?\'\`@+]',' ',x),text2))
    ### for removing the mobile number
    text4=list(map(lambda x: re.sub(r'[0-9]{10}',' ',x),text3))
    text5=list(map(lambda x: re.sub(r'\+[0-9]{2}',' ',x),text4))
    ### https://stackoverflow.com/questions/11331982/how-to-remove-any-url-within-a-string-in-python/11332580
    text6=list(map(lambda x: re.sub(r'\. ',"",x),text5))
    
    review_text = list(map(lambda x: re.sub(r"\'s", " 's ",x),text6))
    review_text = list(map(lambda x: re.sub(r"\'ve", " 've ",x), review_text))
    review_text = list(map(lambda x: re.sub(r"n\'t", " 't ",x), review_text))
    review_text = list(map(lambda x:  re.sub(r"\'re", " 're ",x), review_text))
    review_text = list(map(lambda x: re.sub(r"\'d", " 'd ",x), review_text))
    review_text = list(map(lambda x: re.sub(r"\'ll", " 'll ",x), review_text))
    review_text = list(map(lambda x: re.sub(r",", " ",x), review_text))
    review_text = list(map(lambda x: re.sub(r"!", " ",x), review_text))
    review_text = list(map(lambda x: re.sub(r"\(", "  ",x), review_text))
    review_text = list(map(lambda x: re.sub(r"\)", "  ",x), review_text))
    review_text = list(map(lambda x: re.sub(r"\?", " ",x), review_text))
    review_text = list(map(lambda x: re.sub(r"\s{2,}", " ",x), review_text))
    stemms=[]
    stemmer = SnowballStemmer('english')
    for i in range(0, len(review_text)):
        stemmed_words = [stemmer.stem(word) for word in review_text[i].split()]
        stemms.append(stemmed_words)
    
    
    stops = set(stopwords.words("english"))

    resume_final=[]
    for i in range(len(stemms)):
        words = [w for w in stemms[i] if not w in stops]
        resume_final.append(words)
    return(resume_final)
 

## fiirsr remove email & mobile, website link keep it and then clean data
res_final=processing(text)

#text_process1=list(map(lambda x : re.sub(r"[^A-Za-z0-9(),!.?\'\`]", " ",x),text))
##remove gmail , numbers and names if possible
##take brackets ( as " ( " 



df=pd.DataFrame({"resume":res_final,"entries":file}) ###use zip for more columns

df_1=df.drop(df[df.resume.apply(lambda x : len(x)==0)].index).reset_index()




readjd= open('D:/resumepdf/Net_jd.pdf','rb')
pdfjdReader = PyPDF2.PdfFileReader(readjd)
pagjdobj=pdfjdReader.getPage(pdfjdReader.numPages-1)
jdtext=pagjdobj.extractText()

jdd=processing([jdtext])

jdfinal1=pd.DataFrame({'jdtext':jdd, 'file':'jd'}) ### need to append to combined_tagged for the doc2 vec

### nedd to check for using skip gram oe cbow

###use pca for the averaging the the vectorof the given jd

##use doc to vec as pragraph or doc has unique vector along with the  words in the vector   https://kanoki.org/2019/03/07/sentence-similarity-in-python-using-doc2vec/


from gensim.models.doc2vec import Doc2Vec, TaggedDocument

combined_tagged=[]
for i in range (0,len(df_1)):
    combined_tagged.append(TaggedDocument(df_1.resume[i],df_1.entries[i]))
for i in range(0,len(jdfinal1)):
    combined_tagged.append(TaggedDocument(jdfinal1.jdtext[i],jdfinal1.file[i]))


model = Doc2Vec(dm = 1, min_count=1, window=10, vector_size=150, sample=1e-4, negative=10)
model.build_vocab(combined_tagged)

for epoch in range(15):
     model.train(combined_tagged,total_examples=model.corpus_count, epochs=epoch)
     print("Epoch #{} is complete.".format(epoch+1))
           
words=list(model.wv.vocab)
model.wv.most_similar('.net')

resume_final1=[x for x in res_final if x != []]

pair=[]
for k in range(0, len(resume_final1)):
    distance = model.wv.n_similarity(resume_final1[k],jdd[0])
    pair.append(distance)

data_out=pd.DataFrame({"resume":df_1.resume, "entries":df_1.entries,"distance":pair})

view=data_out.sort_values(by='distance', ascending= False).reset_index(drop=True)[:10]




import shutil
import logging

source=os.getcwd()

dest='D:/destination'

for filename in os.listdir(dest):
    file_path = os.path.join(dest, filename)
    try:
        if os.path.isfile(file_path) or os.path.islink(file_path):
            os.unlink(file_path)
        elif os.path.isdir(file_path):
            shutil.rmtree(file_path)
    except Exception as e:
        print('Failed to delete %s. Reason: %s' % (file_path, e))
        
filelist = view.entries
for f in filelist:
    
        shutil.copy2(source+'\\'+f, dest)

logging.basicConfig(filename='app.log', filemode='w', format='%(name)s- %(levelname)s - %(message)s')

logging.info('directories moved')


####see visulaiztion of the doc2vec using tsne
